# =============================
# ORGANON DATA SOLUTIONS
# core/rendu.py
# =============================

import pandas as pd
import numpy as np
import os
import sys
from datetime import datetime

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import COLORS, PATHS, ORGANON_SIGNATURE, FONTS

# ReportLab
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer,
    Image, Table, TableStyle, PageBreak,
    HRFlowable,
)

# Python-docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Openpyxl
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side


# =============================
# UTILITAIRES CONSOLE
# =============================

def _separateur():
    print("  " + "-" * 50)

def _titre(texte):
    print(f"\n[ {texte} ]")

def _info(label, valeur):
    print(f"  {label:<30}: {valeur}")

def _ok(texte):
    print(f"  {texte} ✓")

def _erreur(texte):
    print(f"  ERREUR — {texte}")


# =============================
# UTILITAIRES STYLES REPORTLAB
# =============================

def _styles_pdf():
    """Retourne les styles Organon pour ReportLab."""
    styles = getSampleStyleSheet()

    style_titre_page = ParagraphStyle(
        "OrganonTitrePage",
        parent=styles["Title"],
        fontSize=24,
        textColor=colors.HexColor(COLORS["dark"]),
        spaceAfter=16,
        fontName="Times-Bold",
        alignment=1,
    )

    style_h1 = ParagraphStyle(
        "OrganonH1",
        parent=styles["Heading1"],
        fontSize=16,
        textColor=colors.HexColor(COLORS["dark"]),
        spaceAfter=10,
        spaceBefore=20,
        fontName="Times-Bold",
    )

    style_h2 = ParagraphStyle(
        "OrganonH2",
        parent=styles["Heading2"],
        fontSize=13,
        textColor=colors.HexColor(COLORS["primary"]),
        spaceAfter=8,
        spaceBefore=12,
        fontName="Times-Bold",
    )

    style_corps = ParagraphStyle(
        "OrganonCorps",
        parent=styles["Normal"],
        fontSize=11,
        textColor=colors.HexColor(COLORS["text"]),
        spaceAfter=8,
        leading=18,
        fontName="Times-Roman",
    )

    style_legende = ParagraphStyle(
        "OrganonLegende",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.HexColor(COLORS["anthracite"]),
        spaceAfter=6,
        alignment=1,
        fontName="Times-Italic",
    )

    style_signature = ParagraphStyle(
        "OrganonSignature",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.HexColor(COLORS["primary"]),
        alignment=2,
        fontName="Times-Italic",
    )

    return {
        "titre_page": style_titre_page,
        "h1":         style_h1,
        "h2":         style_h2,
        "corps":      style_corps,
        "legende":    style_legende,
        "signature":  style_signature,
    }


def _style_tableau_pdf():
    """Style Organon pour les tableaux ReportLab."""
    return TableStyle([
        # Entête
        ("BACKGROUND",    (0, 0), (-1, 0),  colors.HexColor(COLORS["primary"])),
        ("TEXTCOLOR",     (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",      (0, 0), (-1, 0),  "Times-Bold"),
        ("FONTSIZE",      (0, 0), (-1, 0),  11),
        ("ALIGN",         (0, 0), (-1, 0),  "CENTER"),
        # Corps
        ("FONTNAME",      (0, 1), (-1, -1), "Times-Roman"),
        ("FONTSIZE",      (0, 1), (-1, -1), 10),
        ("ROWBACKGROUNDS",(0, 1), (-1, -1),
            [colors.white, colors.HexColor("#F9F7F4")]),
        # Grille
        ("GRID",          (0, 0), (-1, -1), 0.5,
            colors.HexColor("#E5E7EB")),
        ("ALIGN",         (0, 1), (-1, -1), "CENTER"),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("PADDING",       (0, 0), (-1, -1), 7),
    ])


# =============================
# RENDU PDF
# =============================

def rendu_pdf(
    nom_fichier,
    meta,
    rapport_nettoyage,
    resultats_stats,
    fichiers_graphiques,
    dossier=None,
):
    """
    Génère un rapport PDF complet.

    Paramètres :
      nom_fichier         : str  — ex: "rapport_projet1.pdf"
      meta                : dict — titre, auteur, date, contexte,
                                   objectifs, population, source
      rapport_nettoyage   : dict — retourné par nettoyage.py
      resultats_stats     : dict — retourné par stats.py
      fichiers_graphiques : list — chemins PNG retournés par visualisation.py
      dossier             : str  — dossier de sortie (défaut: rapports/)
    """
    if dossier is None:
        dossier = PATHS["rapports"]
    os.makedirs(dossier, exist_ok=True)

    chemin = os.path.join(dossier, nom_fichier)
    styles = _styles_pdf()
    date   = meta.get("date", datetime.today().strftime("%d/%m/%Y"))

    doc = SimpleDocTemplate(
        chemin,
        pagesize=A4,
        rightMargin=2.5*cm,
        leftMargin=2.5*cm,
        topMargin=2.5*cm,
        bottomMargin=2.5*cm,
    )

    contenu = []

    # ---- PAGE DE GARDE ----
    contenu.append(Spacer(1, 2*cm))
    contenu.append(Paragraph(
        meta.get("titre", "Rapport d'analyse"),
        styles["titre_page"]
    ))
    contenu.append(Spacer(1, 0.5*cm))
    contenu.append(HRFlowable(
        width="100%", thickness=2,
        color=colors.HexColor(COLORS["primary"])
    ))
    contenu.append(Spacer(1, 0.5*cm))
    contenu.append(Paragraph(
        f"Auteur : {meta.get('auteur', '')}",
        styles["corps"]
    ))
    contenu.append(Paragraph(
        f"Date : {date}",
        styles["corps"]
    ))
    contenu.append(Spacer(1, 2*cm))
    contenu.append(Paragraph(ORGANON_SIGNATURE, styles["signature"]))
    contenu.append(PageBreak())

    # ---- 1. CONTEXTE ET OBJECTIFS ----
    contenu.append(Paragraph("1. Contexte et objectifs", styles["h1"]))
    contenu.append(HRFlowable(
        width="100%", thickness=1,
        color=colors.HexColor(COLORS["primary"])
    ))
    contenu.append(Spacer(1, 0.3*cm))

    contenu.append(Paragraph("Contexte", styles["h2"]))
    contenu.append(Paragraph(
        meta.get("contexte", "À compléter."),
        styles["corps"]
    ))

    contenu.append(Paragraph("Objectifs", styles["h2"]))
    contenu.append(Paragraph(
        meta.get("objectifs", "À compléter."),
        styles["corps"]
    ))

    contenu.append(Paragraph("Population cible", styles["h2"]))
    contenu.append(Paragraph(
        meta.get("population", "À compléter."),
        styles["corps"]
    ))

    contenu.append(PageBreak())

    # ---- 2. MÉTHODOLOGIE ----
    contenu.append(Paragraph("2. Méthodologie", styles["h1"]))
    contenu.append(HRFlowable(
        width="100%", thickness=1,
        color=colors.HexColor(COLORS["primary"])
    ))
    contenu.append(Spacer(1, 0.3*cm))

    contenu.append(Paragraph("Source des données", styles["h2"]))
    contenu.append(Paragraph(
        meta.get("source", "À compléter."),
        styles["corps"]
    ))

    # Rapport de nettoyage
    if rapport_nettoyage:
        contenu.append(Paragraph("Traitement des données", styles["h2"]))

        lignes_avant = rapport_nettoyage.get("lignes_avant", "—")
        lignes_apres = rapport_nettoyage.get("lignes_apres", "—")
        doublons     = rapport_nettoyage.get("doublons_supprimes", 0)
        vides        = rapport_nettoyage.get("lignes_vides_supprimees", 0)
        manquantes   = rapport_nettoyage.get("valeurs_manquantes", {})

        data_nett = [
            ["Indicateur", "Valeur"],
            ["Lignes avant nettoyage",      str(lignes_avant)],
            ["Lignes après nettoyage",      str(lignes_apres)],
            ["Doublons supprimés",           str(doublons)],
            ["Lignes vides supprimées",      str(vides)],
            ["Variables avec manquants",     str(len(manquantes))],
        ]

        table = Table(data_nett, colWidths=[10*cm, 5*cm])
        table.setStyle(_style_tableau_pdf())
        contenu.append(table)
        contenu.append(Spacer(1, 0.5*cm))

    contenu.append(Paragraph("Outils utilisés", styles["h2"]))
    contenu.append(Paragraph(
        "Analyse réalisée avec Python 3 — bibliothèques : "
        "pandas, numpy, matplotlib, seaborn. "
        "Rapport généré par Organon Data Solutions.",
        styles["corps"]
    ))

    contenu.append(PageBreak())

    # ---- 3. RÉSULTATS ----
    contenu.append(Paragraph("3. Résultats", styles["h1"]))
    contenu.append(HRFlowable(
        width="100%", thickness=1,
        color=colors.HexColor(COLORS["primary"])
    ))
    contenu.append(Spacer(1, 0.3*cm))

    # Tableau stats numériques
    if resultats_stats and resultats_stats.get("numeriques"):
        contenu.append(Paragraph(
            "3.1 Statistiques descriptives — Variables numériques",
            styles["h2"]
        ))

        entetes = [
            "Variable", "N", "Moyenne", "Médiane",
            "Éc.-type", "Min", "Max", "Outliers"
        ]
        data_num = [entetes]

        for col, r in resultats_stats["numeriques"].items():
            data_num.append([
                col,
                str(r["n"]),
                str(r["moyenne"]),
                str(r["mediane"]),
                str(r["ecart_type"]),
                str(r["min"]),
                str(r["max"]),
                str(r["outliers"]),
            ])

        table_num = Table(data_num, repeatRows=1)
        table_num.setStyle(_style_tableau_pdf())
        contenu.append(table_num)
        contenu.append(Spacer(1, 0.5*cm))

    # Tableau stats catégorielles
    if resultats_stats and resultats_stats.get("categories"):
        contenu.append(Paragraph(
            "3.2 Statistiques descriptives — Variables catégorielles",
            styles["h2"]
        ))

        for col, r in resultats_stats["categories"].items():
            contenu.append(Paragraph(f"Variable : {col}", styles["corps"]))

            entetes_cat = ["Modalité", "Effectif", "Pourcentage"]
            data_cat    = [entetes_cat]

            for modalite, vals in r["frequences"].items():
                data_cat.append([
                    str(modalite),
                    str(vals["n"]),
                    f"{vals['pct']}%",
                ])

            table_cat = Table(data_cat, repeatRows=1, colWidths=[7*cm, 4*cm, 4*cm])
            table_cat.setStyle(_style_tableau_pdf())
            contenu.append(table_cat)
            contenu.append(Spacer(1, 0.4*cm))

    contenu.append(PageBreak())

    # Graphiques
    if fichiers_graphiques:
        contenu.append(Paragraph("3.3 Visualisations", styles["h2"]))
        contenu.append(Spacer(1, 0.3*cm))

        for chemin_graph in fichiers_graphiques:
            if os.path.exists(chemin_graph):
                nom_graph = os.path.basename(chemin_graph)
                try:
                    contenu.append(Image(chemin_graph, width=15*cm, height=9*cm))
                    contenu.append(Paragraph(
                        f"Figure : {nom_graph}",
                        styles["legende"]
                    ))
                    contenu.append(Spacer(1, 0.5*cm))
                except Exception as e:
                    _erreur(f"Image non chargée : {nom_graph} — {e}")
            else:
                _erreur(f"Fichier introuvable : {chemin_graph}")

    contenu.append(PageBreak())

    # ---- 4. INTERPRÉTATION ----
    contenu.append(Paragraph("4. Interprétation", styles["h1"]))
    contenu.append(HRFlowable(
        width="100%", thickness=1,
        color=colors.HexColor(COLORS["primary"])
    ))
    contenu.append(Spacer(1, 0.3*cm))

    # Interprétation automatique depuis stats
    if resultats_stats and resultats_stats.get("numeriques"):
        for col, r in resultats_stats["numeriques"].items():
            contenu.append(Paragraph(col, styles["h2"]))

            if abs(r["asymetrie"]) < 0.5:
                interp = "La distribution est approximativement symétrique."
            elif r["asymetrie"] > 0.5:
                interp = (
                    "La distribution présente une asymétrie positive "
                    "(queue longue vers les grandes valeurs)."
                )
            else:
                interp = (
                    "La distribution présente une asymétrie négative "
                    "(queue longue vers les petites valeurs)."
                )

            if r["outliers"] > 0:
                interp += (
                    f" {r['outliers']} valeur(s) aberrante(s) ont été détectées "
                    "et méritent une vérification."
                )

            contenu.append(Paragraph(interp, styles["corps"]))

    contenu.append(Paragraph(
        meta.get("interpretation", ""),
        styles["corps"]
    ))

    contenu.append(PageBreak())

    # ---- 5. ANNEXES ----
    contenu.append(Paragraph("5. Annexes", styles["h1"]))
    contenu.append(HRFlowable(
        width="100%", thickness=1,
        color=colors.HexColor(COLORS["primary"])
    ))
    contenu.append(Spacer(1, 0.3*cm))
    contenu.append(Paragraph(
        "Les données complètes nettoyées sont disponibles "
        "dans le fichier Excel joint.",
        styles["corps"]
    ))

    # Pied de page final
    contenu.append(Spacer(1, 2*cm))
    contenu.append(HRFlowable(
        width="100%", thickness=1,
        color=colors.HexColor(COLORS["primary"])
    ))
    contenu.append(Spacer(1, 0.3*cm))
    contenu.append(Paragraph(ORGANON_SIGNATURE, styles["signature"]))
    contenu.append(Paragraph(
        f"Rapport généré le {date}",
        styles["signature"]
    ))

    # Construction
    try:
        doc.build(contenu)
        _ok(f"PDF exporté → {chemin}")
        return chemin
    except Exception as e:
        _erreur(f"PDF : {e}")
        return None


# =============================
# RENDU EXCEL
# =============================

def rendu_excel(nom_fichier, df, resultats_stats, dossier=None):
    """
    Génère un fichier Excel avec :
    - Feuille 1 : données nettoyées
    - Feuille 2 : statistiques numériques
    - Feuille 3 : statistiques catégorielles
    """
    if dossier is None:
        dossier = PATHS["rapports"]
    os.makedirs(dossier, exist_ok=True)

    chemin = os.path.join(dossier, nom_fichier)

    try:
        with pd.ExcelWriter(chemin, engine="openpyxl") as writer:

            # Feuille 1 — Données
            df.to_excel(writer, sheet_name="Données", index=False)
            ws1 = writer.sheets["Données"]
            _formater_entete_excel(ws1)
            _ajuster_colonnes_excel(ws1)

            # Feuille 2 — Stats numériques
            if resultats_stats and resultats_stats.get("numeriques"):
                rows_num = []
                for col, r in resultats_stats["numeriques"].items():
                    rows_num.append({
                        "Variable":     col,
                        "N":            r["n"],
                        "Moyenne":      r["moyenne"],
                        "Médiane":      r["mediane"],
                        "Écart-type":   r["ecart_type"],
                        "Min":          r["min"],
                        "Max":          r["max"],
                        "Q1":           r["q1"],
                        "Q3":           r["q3"],
                        "IQR":          r["iqr"],
                        "Asymétrie":    r["asymetrie"],
                        "Aplatissement":r["aplatiss"],
                        "CV (%)":       r["cv"],
                        "Outliers":     r["outliers"],
                    })
                df_num = pd.DataFrame(rows_num)
                df_num.to_excel(
                    writer, sheet_name="Stats numériques", index=False
                )
                ws2 = writer.sheets["Stats numériques"]
                _formater_entete_excel(ws2)
                _ajuster_colonnes_excel(ws2)

            # Feuille 3 — Stats catégorielles
            if resultats_stats and resultats_stats.get("categories"):
                rows_cat = []
                for col, r in resultats_stats["categories"].items():
                    for modalite, vals in r["frequences"].items():
                        rows_cat.append({
                            "Variable":   col,
                            "Modalité":   modalite,
                            "Effectif":   vals["n"],
                            "Pourcentage": f"{vals['pct']}%",
                        })
                df_cat = pd.DataFrame(rows_cat)
                df_cat.to_excel(
                    writer, sheet_name="Stats catégorielles", index=False
                )
                ws3 = writer.sheets["Stats catégorielles"]
                _formater_entete_excel(ws3)
                _ajuster_colonnes_excel(ws3)

        _ok(f"Excel exporté → {chemin}")
        return chemin

    except Exception as e:
        _erreur(f"Excel : {e}")
        return None


def _formater_entete_excel(ws):
    """Applique le style Organon à l'entête d'une feuille Excel."""
    fill   = PatternFill("solid", fgColor="CD7F32")
    font   = Font(bold=True, color="FFFFFF", name="Times New Roman")
    align  = Alignment(horizontal="center", vertical="center")
    for cell in ws[1]:
        cell.fill      = fill
        cell.font      = font
        cell.alignment = align


def _ajuster_colonnes_excel(ws):
    """Ajuste la largeur des colonnes automatiquement."""
    for col in ws.columns:
        max_len = max(
            len(str(cell.value or ""))
            for cell in col
        )
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)


# =============================
# RENDU WORD
# =============================

def rendu_word(
    nom_fichier,
    meta,
    rapport_nettoyage,
    resultats_stats,
    fichiers_graphiques,
    dossier=None,
):
    """
    Génère un rapport Word éditable.
    Même structure que le PDF.
    """
    if dossier is None:
        dossier = PATHS["rapports"]
    os.makedirs(dossier, exist_ok=True)

    chemin = os.path.join(dossier, nom_fichier)
    date   = meta.get("date", datetime.today().strftime("%d/%m/%Y"))

    try:
        doc = Document()

        # Style global
        style_normal = doc.styles["Normal"]
        style_normal.font.name = "Times New Roman"
        style_normal.font.size = Pt(11)

        # ---- PAGE DE GARDE ----
        titre = doc.add_heading(meta.get("titre", "Rapport d'analyse"), 0)
        titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titre.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        doc.add_paragraph(f"Auteur : {meta.get('auteur', '')}")
        doc.add_paragraph(f"Date   : {date}")
        doc.add_paragraph(ORGANON_SIGNATURE)
        doc.add_page_break()

        # ---- 1. CONTEXTE ----
        doc.add_heading("1. Contexte et objectifs", 1)
        doc.add_heading("Contexte", 2)
        doc.add_paragraph(meta.get("contexte", "À compléter."))
        doc.add_heading("Objectifs", 2)
        doc.add_paragraph(meta.get("objectifs", "À compléter."))
        doc.add_heading("Population cible", 2)
        doc.add_paragraph(meta.get("population", "À compléter."))
        doc.add_page_break()

        # ---- 2. MÉTHODOLOGIE ----
        doc.add_heading("2. Méthodologie", 1)
        doc.add_heading("Source des données", 2)
        doc.add_paragraph(meta.get("source", "À compléter."))

        if rapport_nettoyage:
            doc.add_heading("Traitement des données", 2)
            table_nett = doc.add_table(rows=1, cols=2)
            table_nett.style = "Table Grid"
            hdr = table_nett.rows[0].cells
            hdr[0].text = "Indicateur"
            hdr[1].text = "Valeur"
            for label, val in [
                ("Lignes avant", rapport_nettoyage.get("lignes_avant", "—")),
                ("Lignes après", rapport_nettoyage.get("lignes_apres", "—")),
                ("Doublons supprimés", rapport_nettoyage.get("doublons_supprimes", 0)),
                ("Lignes vides supprimées", rapport_nettoyage.get("lignes_vides_supprimees", 0)),
            ]:
                row = table_nett.add_row().cells
                row[0].text = str(label)
                row[1].text = str(val)

        doc.add_page_break()

        # ---- 3. RÉSULTATS ----
        doc.add_heading("3. Résultats", 1)

        # Stats numériques
        if resultats_stats and resultats_stats.get("numeriques"):
            doc.add_heading(
                "3.1 Statistiques descriptives — Variables numériques", 2
            )
            entetes = [
                "Variable", "N", "Moyenne", "Médiane",
                "Éc.-type", "Min", "Max", "Outliers"
            ]
            table_num = doc.add_table(rows=1, cols=len(entetes))
            table_num.style = "Table Grid"
            for i, h in enumerate(entetes):
                table_num.rows[0].cells[i].text = h

            for col, r in resultats_stats["numeriques"].items():
                row = table_num.add_row().cells
                vals = [
                    col, str(r["n"]), str(r["moyenne"]),
                    str(r["mediane"]), str(r["ecart_type"]),
                    str(r["min"]), str(r["max"]), str(r["outliers"]),
                ]
                for i, v in enumerate(vals):
                    row[i].text = v

        # Stats catégorielles
        if resultats_stats and resultats_stats.get("categories"):
            doc.add_heading(
                "3.2 Statistiques descriptives — Variables catégorielles", 2
            )
            for col, r in resultats_stats["categories"].items():
                doc.add_paragraph(f"Variable : {col}")
                table_cat = doc.add_table(rows=1, cols=3)
                table_cat.style = "Table Grid"
                hdr = table_cat.rows[0].cells
                hdr[0].text = "Modalité"
                hdr[1].text = "Effectif"
                hdr[2].text = "Pourcentage"
                for modalite, vals in r["frequences"].items():
                    row = table_cat.add_row().cells
                    row[0].text = str(modalite)
                    row[1].text = str(vals["n"])
                    row[2].text = f"{vals['pct']}%"

        doc.add_page_break()

        # Graphiques
        if fichiers_graphiques:
            doc.add_heading("3.3 Visualisations", 2)
            for chemin_graph in fichiers_graphiques:
                if os.path.exists(chemin_graph):
                    try:
                        doc.add_picture(chemin_graph, width=Inches(5.5))
                        doc.add_paragraph(
                            os.path.basename(chemin_graph)
                        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
                    except Exception as e:
                        _erreur(f"Image non insérée : {e}")
                else:
                    _erreur(f"Fichier introuvable : {chemin_graph}")

        doc.add_page_break()

        # ---- 4. INTERPRÉTATION ----
        doc.add_heading("4. Interprétation", 1)

        if resultats_stats and resultats_stats.get("numeriques"):
            for col, r in resultats_stats["numeriques"].items():
                doc.add_heading(col, 2)
                if abs(r["asymetrie"]) < 0.5:
                    interp = "Distribution approximativement symétrique."
                elif r["asymetrie"] > 0.5:
                    interp = (
                        "Distribution asymétrique positive "
                        "(queue longue vers les grandes valeurs)."
                    )
                else:
                    interp = (
                        "Distribution asymétrique négative "
                        "(queue longue vers les petites valeurs)."
                    )
                if r["outliers"] > 0:
                    interp += (
                        f" {r['outliers']} valeur(s) aberrante(s) détectée(s)."
                    )
                doc.add_paragraph(interp)

        if meta.get("interpretation"):
            doc.add_paragraph(meta["interpretation"])

        doc.add_page_break()

        # ---- 5. ANNEXES ----
        doc.add_heading("5. Annexes", 1)
        doc.add_paragraph(
            "Les données complètes nettoyées sont disponibles "
            "dans le fichier Excel joint."
        )

        # Pied de page
        doc.add_paragraph()
        doc.add_paragraph(ORGANON_SIGNATURE)
        doc.add_paragraph(f"Rapport généré le {date}")

        doc.save(chemin)
        _ok(f"Word exporté → {chemin}")
        return chemin

    except Exception as e:
        _erreur(f"Word : {e}")
        return None


# =============================
# PIPELINE PRINCIPAL
# =============================

def pipeline_rendu(
    nom_projet,
    meta,
    df,
    rapport_nettoyage,
    resultats_stats,
    fichiers_graphiques,
    dossier=None,
):
    """
    Pipeline interactif de rendu.
    Laisse l'utilisateur choisir le format.

    Usage :
      from core.rendu import pipeline_rendu
      pipeline_rendu(
          nom_projet        = "projet1",
          meta              = {...},
          df                = df,
          rapport_nettoyage = rapport,
          resultats_stats   = resultats,
          fichiers_graphiques = fichiers,
      )
    """
    _titre("RENDU")
    print("  Formats disponibles :")
    print("    1. PDF        — rapport complet")
    print("    2. Excel      — données + statistiques")
    print("    3. Word       — rapport éditable")
    print("    4. Tout       — les trois formats")
    print()

    choix_raw = input("  Votre choix (ex: 1,2) : ").strip()

    if not choix_raw:
        print("  Aucun rendu produit.")
        return []

    # Option 4 = tout
    if "4" in choix_raw:
        choix = ["1", "2", "3"]
    else:
        choix = [c.strip() for c in choix_raw.split(",") if c.strip() in ["1","2","3"]]

    fichiers_produits = []

    for c in choix:
        if c == "1":
            print(f"\n  PDF   : en cours...")
            chemin = rendu_pdf(
                f"rapport_{nom_projet}.pdf",
                meta,
                rapport_nettoyage,
                resultats_stats,
                fichiers_graphiques,
                dossier,
            )
            if chemin:
                fichiers_produits.append(chemin)

        elif c == "2":
            print(f"\n  Excel : en cours...")
            chemin = rendu_excel(
                f"{nom_projet}_donnees.xlsx",
                df,
                resultats_stats,
                dossier,
            )
            if chemin:
                fichiers_produits.append(chemin)

        elif c == "3":
            print(f"\n  Word  : en cours...")
            chemin = rendu_word(
                f"rapport_{nom_projet}.docx",
                meta,
                rapport_nettoyage,
                resultats_stats,
                fichiers_graphiques,
                dossier,
            )
            if chemin:
                fichiers_produits.append(chemin)

    print()
    print("=" * 52)
    print(f"  [ RENDU ] Terminé ✓")
    print(f"  Fichiers produits : {len(fichiers_produits)}")
    for f in fichiers_produits:
        print(f"    — {f}")
    print("=" * 52)
    print()

    return fichiers_produits